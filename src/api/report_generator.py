# ==========================================================
# report_generator.py
# Populates the Sample.docx template with approved RCA/CAPA
# content and user-provided metadata.
#
# Strategy: open the template, modify specific cells/paragraphs
# in-place, save to a new file. This preserves all formatting,
# fonts, colours, borders and styles exactly.
#
# What gets populated:
#   - Header: shortened problem description title
#   - Table 3 (Technical Investigation): cleared + replaced
#   - Table 4 (5-Why): cleared + replaced with AI content
#   - Table 5 (Root Cause Summary): AI content, N/A for category
#   - Table 6 (CAPA): cleared + replaced with AI content
#   - Para 58 (Lessons Learned): replaced with AI content
#   - Table 1 rows 1,2,4,5,6,7: user input fields
#
# What is left unchanged:
#   - Table 0 (Approvers) — user fills after download
#   - Table 1 row 0 (Problem Number) — user fills after download
#   - Table 1 row 3 (Patient Safety) — left as-is
#   - Table 2 (Related Tickets) — left as-is, user fills
#   - Table 5 col 1 (Root cause category) — left as N/A
#   - Table 6 cols 3,4,5 (Due date, Status, Ticket) — user fills
#   - Table 7 (Version History) — left as-is
#   - All notes, TOC, headings, body text — left as-is
# ==========================================================

import copy
import os
import logging
import subprocess
from typing import Any, Dict, List

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Paths ─────────────────────────────────────────────────
TEMPLATE_PATH = "/content/rca-slm/templates/Sample.docx"
OUTPUT_DIR    = "/content/rca-slm/outputs"


# ── Cell text helpers ─────────────────────────────────────

def _clear_cell(cell) -> None:
    """
    Removes all paragraphs from a cell except the first one,
    then clears the first paragraph's text.
    Preserves cell formatting (borders, shading, width).
    """
    # Remove all paragraphs after the first
    for para in cell.paragraphs[1:]:
        p = para._element
        p.getparent().remove(p)
    # Clear text in the first paragraph while keeping its runs' formatting
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
        # If no runs, clear directly
        if not para.runs:
            para.clear()


def _set_cell_text(cell, text: str) -> None:
    """
    Sets plain text in a cell, preserving the cell's existing
    paragraph style and run formatting (font, size, colour).

    Clears existing content first, then writes new text into
    the first run of the first paragraph to inherit formatting.
    """
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return

    para = cell.paragraphs[0]

    # Remove extra paragraphs
    for extra in cell.paragraphs[1:]:
        p = extra._element
        p.getparent().remove(p)

    if para.runs:
        # Write into first run to inherit font/size/colour
        para.runs[0].text = text
        # Clear any remaining runs
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _clear_content_controls(cell) -> None:
    """
    Removes Word content controls (sdt elements) from a cell.
    These appear as dropdown fields like 'Choose an item.' in the template.
    After removal, the cell is treated as plain text.
    """
    # Find and remove all sdt (structured document tag) elements
    for sdt in cell._element.findall('.//' + qn('w:sdt')):
        sdt.getparent().remove(sdt)


def _copy_row_format(source_row, target_row) -> None:
    """
    Copies row-level XML formatting (height, style) from source to target.
    Used when adding new rows to match template row appearance.
    """
    source_tr = source_row._tr
    target_tr = target_row._tr
    source_trPr = source_tr.find(qn('w:trPr'))
    if source_trPr is not None:
        target_trPr = copy.deepcopy(source_trPr)
        existing = target_tr.find(qn('w:trPr'))
        if existing is not None:
            target_tr.remove(existing)
        target_tr.insert(0, target_trPr)


def _add_table_row(table, template_row, values: List[str]) -> None:
    """
    Adds a new row to a table by copying the template_row's XML structure
    (preserving cell widths, borders, font formatting) and filling in values.

    Args:
        table:        The docx Table object to add the row to.
        template_row: An existing row whose format to copy.
        values:       List of strings, one per cell in the row.
    """
    # Deep copy the template row's XML
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)

    # Get the newly added row object
    new_row = table.rows[-1]

    # Fill in the cell values
    for i, value in enumerate(values):
        if i < len(new_row.cells):
            _clear_content_controls(new_row.cells[i])
            _set_cell_text(new_row.cells[i], value)


# ── Header population ─────────────────────────────────────

def _populate_header(doc: Document, problem_description: str) -> None:
    """
    Updates the document header table with a shortened problem description.
    The header appears on every page — updating it once updates all pages.

    Cell 0: shortened incident title (max 80 chars)
    Cell 1: left as-is (PRB00xxxx regarding INCxxxxxx — user fills)
    """
    for section in doc.sections:
        header = section.header
        if header.tables:
            header_table = header.tables[0]
            if header_table.rows and len(header_table.rows[0].cells) >= 1:
                # Use first 120 chars of problem description as the title
                short_title = f"IT RCA — {problem_description[:120]}"
                if len(problem_description) > 120:
                    short_title += "..."
                _set_cell_text(header_table.rows[0].cells[0], short_title)
                # Cell 1 (problem number) — left unchanged, user fills


# ── Table 1 — Problem Summary ─────────────────────────────
def _populate_problem_summary(
    doc: Document,
    problem_description: str,
    business_impact: str,
    impacted_location: str,
    impacted_business_unit: str,
    impacted_service_tower: str,
    impacted_application: str,
) -> None:
    """
    Populates Table 1 (Problem Summary) with user-provided incident details.

    Row 0 (Problem Number) — left as-is, user fills after download.
    Row 3 (Patient Safety) — left as-is.
    Row 5 (Business Unit) — has only 1 cell in the template XML.
                            We append the value directly into cell 0
                            after the label text using a line break.
    """
    table = doc.tables[1]

    # Row 1 — Problem Description
    _clear_content_controls(table.rows[1].cells[1])
    _set_cell_text(table.rows[1].cells[1], problem_description)

    # Row 2 — Business Impact
    _clear_content_controls(table.rows[2].cells[1])
    _set_cell_text(table.rows[2].cells[1], business_impact)

    # Row 3 — Patient Safety — leave as-is

    # Row 4 — Impacted Location
    _clear_content_controls(table.rows[4].cells[1])
    _set_cell_text(table.rows[4].cells[1], impacted_location)

    # Row 5 — Impacted Business Unit (only 1 cell — no value cell in template)
    # Add the value as a new paragraph inside the single cell,
    # separated from the label so it reads clearly
    cell5 = table.rows[5].cells[0]
    _clear_content_controls(cell5)
    # Keep existing label text, add value as new paragraph
    para = cell5.add_paragraph()
    para.add_run(impacted_business_unit)

    # Row 6 — Impacted Service Tower
    _clear_content_controls(table.rows[6].cells[1])
    _set_cell_text(table.rows[6].cells[1], impacted_service_tower)

    # Row 7 — Impacted Application
    _clear_content_controls(table.rows[7].cells[1])
    _set_cell_text(table.rows[7].cells[1], impacted_application)

    

# ── Table 3 — Technical Investigation ────────────────────

def _populate_technical_investigation(
    doc: Document,
    timeline: List[Dict[str, str]],
) -> None:
    """
    Clears all existing data rows from Table 3 (Technical Investigation)
    and replaces them with user-provided timeline entries.

    Each timeline entry: {date, time, activity}
    Rows 0 and 1 are the merged header and column header rows — kept as-is.
    Rows 2+ are data rows — all cleared, then replaced.

    Args:
        timeline: List of dicts with keys: date, time, activity.
    """
    table = doc.tables[3]

    # Remove all existing data rows (rows 2 onwards)
    # Must remove from the end to avoid index shifting
    data_rows = table.rows[2:]
    for row in data_rows:
        row._element.getparent().remove(row._element)

    if not timeline:
        return

    # Use the original row 2 as format template (already removed from table)
    # Instead use row 1 (column header) as format reference
    template_row = table.rows[1]

    for entry in timeline:
        _add_table_row(
            table=table,
            template_row=template_row,
            values=[
                entry.get("date", ""),
                entry.get("time", ""),
                entry.get("activity", ""),
            ]
        )


# ── Table 4 — 5-Why Analysis ──────────────────────────────
def _populate_five_why(
    doc: Document,
    five_why_analysis: List[Dict[str, str]],
) -> None:
    """
    Clears all existing data rows from Table 4 (5-Why Analysis)
    and replaces them with AI-generated why entries.

    Row 0 is the header row — kept as-is.
    Rows 1+ are data rows — cleared and replaced.
    Template row is captured from Row 1 (first data row)
    BEFORE deletion so new rows inherit data row formatting
    (no purple background) not header row formatting.
    """
    table = doc.tables[4]

    # Capture data row format BEFORE deleting — Row 1 is first data row
    # Using Row 0 (header) would copy purple background to all new rows
    import copy
    template_row_xml = copy.deepcopy(table.rows[1]._tr)

    # Remove existing data rows (row 1 onwards)
    data_rows = table.rows[1:]
    for row in data_rows:
        row._element.getparent().remove(row._element)

    # Create a temporary row object from the saved XML to use as template
    from docx.oxml import OxmlElement
    for i, entry in enumerate(five_why_analysis):
        question = entry.get("question", f"Why {i+1}?")
        answer   = entry.get("answer", "")

        # Deep copy the saved data row XML for each new row
        new_tr = copy.deepcopy(template_row_xml)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]

        _clear_content_controls(new_row.cells[0])
        _clear_content_controls(new_row.cells[1])
        _set_cell_text(new_row.cells[0], question)
        _set_cell_text(new_row.cells[1], answer)

# ── Table 5 — Root Cause Summary ─────────────────────────

def _populate_root_cause(
    doc: Document,
    root_cause_summary: Dict[str, str],
) -> None:
    """
    Populates Table 5 Row 1 with the AI-generated root cause statement.

    Cell 0: root cause statement
    Cell 1: left as N/A (per requirement)
    Cell 2: left as-is
    """
    table = doc.tables[5]

    if len(table.rows) > 1:
        _clear_content_controls(table.rows[1].cells[0])
        _set_cell_text(
            table.rows[1].cells[0],
            root_cause_summary.get("statement", "")
        )
        # Cell 1 (category) — leave as N/A per requirement
        # Cell 2 — leave as-is


# ── Table 6 — CAPA ───────────────────────────────────────

def _populate_capa(
    doc: Document,
    corrective_preventive_actions: List[Dict[str, str]],
) -> None:
    """
    Clears all existing data rows from Table 6 (CAPA)
    and replaces them with AI-generated actions.

    Row 0 is the header row — kept as-is.
    Template row captured from Row 1 (first data row) BEFORE
    deletion so new rows have no purple background.
    Columns 3,4,5 (Due Date, Status, Ticket) left blank.
    """
    table = doc.tables[6]

    # Capture data row format BEFORE deleting
    import copy
    template_row_xml = copy.deepcopy(table.rows[1]._tr)

    # Remove existing data rows
    data_rows = table.rows[1:]
    for row in data_rows:
        row._element.getparent().remove(row._element)

    for action in corrective_preventive_actions:
        new_tr = copy.deepcopy(template_row_xml)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]

        # Clear all cells first
        for cell in new_row.cells:
            _clear_content_controls(cell)

        _set_cell_text(new_row.cells[0], action.get("action_type", ""))
        _set_cell_text(new_row.cells[1], action.get("action_description", ""))
        _set_cell_text(new_row.cells[2], action.get("action_owner", ""))
        _set_cell_text(new_row.cells[3], "")   # Due Date — user fills
        _set_cell_text(new_row.cells[4], "")   # Status — user fills
        _set_cell_text(new_row.cells[5], "")   # Ticket — user fills



# ── Paragraph 58 — Lessons Learned ───────────────────────

def _populate_lessons_learned(
    doc: Document,
    lessons_learned: List[str],
) -> None:
    """
    Replaces the content of paragraph 58 (Lessons Learned body text)
    with AI-generated lessons.

    If multiple lessons exist, each is written as a separate paragraph
    inserted after paragraph 58, preserving the Normal style.
    """
    # Find paragraph 58
    if len(doc.paragraphs) <= 58:
        logger.warning("Paragraph 58 not found — lessons learned not populated.")
        return

    lessons_para = doc.paragraphs[58]

    if not lessons_learned:
        return

    # Set first lesson in para 58 to preserve its position in the document
    lessons_para.clear()
    lessons_para.add_run(lessons_learned[0])

    # Insert additional lessons as new paragraphs after para 58
    # using XML insertion to maintain document structure
    parent = lessons_para._element.getparent()
    insert_after = lessons_para._element

    for lesson in lessons_learned[1:]:
        # Create a new paragraph element copying the style of para 58
        new_para = OxmlElement('w:p')
        new_pPr  = copy.deepcopy(lessons_para._element.find(qn('w:pPr')))
        if new_pPr is not None:
            new_para.append(new_pPr)
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = lesson
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_para.append(new_r)
        insert_after.addnext(new_para)
        insert_after = new_para


# ── PDF conversion ────────────────────────────────────────

def _convert_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    Converts a DOCX file to PDF using LibreOffice headless mode.
    LibreOffice is pre-installed in Colab.

    Args:
        docx_path:  Full path to the .docx file.
        output_dir: Directory where the PDF will be saved.

    Returns:
        Full path to the generated PDF file.

    Raises:
        RuntimeError if LibreOffice conversion fails.
    """
    result = subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            docx_path,
            "--outdir", output_dir,
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice PDF conversion failed: {result.stderr}"
        )

    # PDF filename matches DOCX filename with .pdf extension
    pdf_filename = os.path.basename(docx_path).replace(".docx", ".pdf")
    pdf_path = os.path.join(output_dir, pdf_filename)

    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF not found after conversion: {pdf_path}")

    return pdf_path


# ── Main entry point ──────────────────────────────────────

def generate_report(report_data: Dict[str, Any]) -> Dict[str, str]:
    """
    Generates a populated DOCX report and converts it to PDF.

    Opens the Sample.docx template, populates all sections with
    approved AI-generated content and user-provided metadata,
    saves the DOCX, then converts to PDF using LibreOffice.

    Args:
        report_data: Dict containing all approved content and metadata.
            Required keys:
                problem_description, business_impact,
                impacted_location, impacted_business_unit,
                impacted_service_tower, impacted_application,
                technical_investigation (list of {date, time, activity}),
                five_why_analysis (list of {question, answer}),
                root_cause_summary ({statement, root_cause_category}),
                corrective_preventive_actions (list of {action_type,
                    action_description, action_owner}),
                lessons_learned (list of strings)

    Returns:
        Dict with keys:
            docx_path: Full path to the generated DOCX file.
            pdf_path:  Full path to the generated PDF file.

    Raises:
        FileNotFoundError if template is missing.
        RuntimeError if PDF conversion fails.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Load template — always work from a fresh copy to avoid
    # accumulating changes across multiple report generations
    doc = Document(TEMPLATE_PATH)

    # ── Populate each section ──────────────────────────────

    _populate_header(
        doc=doc,
        problem_description=report_data.get("problem_description", ""),
    )

    _populate_problem_summary(
        doc=doc,
        problem_description=report_data.get("problem_description", ""),
        business_impact=report_data.get("business_impact", ""),
        impacted_location=report_data.get("impacted_location", ""),
        impacted_business_unit=report_data.get("impacted_business_unit", ""),
        impacted_service_tower=report_data.get("impacted_service_tower", ""),
        impacted_application=report_data.get("impacted_application", ""),
    )

    _populate_technical_investigation(
        doc=doc,
        timeline=report_data.get("technical_investigation", []),
    )

    _populate_five_why(
        doc=doc,
        five_why_analysis=report_data.get("five_why_analysis", []),
    )

    _populate_root_cause(
        doc=doc,
        root_cause_summary=report_data.get("root_cause_summary", {}),
    )

    _populate_capa(
        doc=doc,
        corrective_preventive_actions=report_data.get(
            "corrective_preventive_actions", []
        ),
    )

    _populate_lessons_learned(
        doc=doc,
        lessons_learned=report_data.get("lessons_learned", []),
    )

    # ── Save DOCX ──────────────────────────────────────────
    docx_filename = "rca_report.docx"
    docx_path     = os.path.join(OUTPUT_DIR, docx_filename)
    doc.save(docx_path)
    logger.info(f"DOCX saved: {docx_path}")

    # ── Convert to PDF ─────────────────────────────────────
    try:
        pdf_path = _convert_to_pdf(docx_path, OUTPUT_DIR)
        logger.info(f"PDF saved: {pdf_path}")
    except RuntimeError as e:
        logger.error(f"PDF conversion failed: {e}")
        pdf_path = ""

    return {
        "docx_path": docx_path,
        "pdf_path":  pdf_path,
    }